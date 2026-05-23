"""
Document loader — supports PDF, TXT, DOCX, Markdown, and web URLs.
Returns normalized list of (text, metadata) tuples for downstream chunking.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

import requests
from bs4 import BeautifulSoup
from loguru import logger


@dataclass
class RawDocument:
    text: str
    metadata: dict = field(default_factory=dict)


def load_file(file_path: str) -> List[RawDocument]:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path)
    elif suffix in (".txt", ".md"):
        return _load_text(path)
    elif suffix == ".docx":
        return _load_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def load_url(url: str) -> List[RawDocument]:
    logger.info(f"Fetching URL: {url}")
    response = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove noise elements
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    text = _clean_whitespace(text)

    title = soup.find("title")
    return [RawDocument(
        text=text,
        metadata={"source": url, "title": title.get_text() if title else url, "type": "url"}
    )]


def _load_pdf(path: Path) -> List[RawDocument]:
    from pypdf import PdfReader

    logger.info(f"Loading PDF: {path.name}")
    reader = PdfReader(str(path))
    docs = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = _clean_whitespace(text)
        if text.strip():
            docs.append(RawDocument(
                text=text,
                metadata={"source": path.name, "page": i + 1, "type": "pdf"}
            ))

    logger.info(f"Loaded {len(docs)} pages from {path.name}")
    return docs


def _load_text(path: Path) -> List[RawDocument]:
    logger.info(f"Loading text file: {path.name}")
    text = path.read_text(encoding="utf-8")
    return [RawDocument(
        text=_clean_whitespace(text),
        metadata={"source": path.name, "type": "text"}
    )]


def _load_docx(path: Path) -> List[RawDocument]:
    from docx import Document

    logger.info(f"Loading DOCX: {path.name}")
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return [RawDocument(
        text=_clean_whitespace(text),
        metadata={"source": path.name, "type": "docx"}
    )]


def _clean_whitespace(text: str) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()
